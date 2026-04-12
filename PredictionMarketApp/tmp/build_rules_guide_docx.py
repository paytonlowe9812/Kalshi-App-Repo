from pathlib import Path
import re
from urllib.parse import urlparse, unquote

from docx import Document
from docx.shared import Inches


MD_PATH = Path(r"c:\Users\pman\Kalshi-App-Repo\PredictionMarketApp\tmp\ruleset-builder-complete-guide.md")
OUT_PATH = Path(r"c:\Users\pman\Kalshi-App-Repo\PredictionMarketApp\tmp\ruleset-builder-complete-guide.docx")


def markdown_image_path(raw_path: str) -> Path:
    raw_path = raw_path.strip()
    if raw_path.lower().startswith("file:///"):
        parsed = urlparse(raw_path)
        local_path = unquote(parsed.path.lstrip("/"))
        return Path(local_path)
    return Path(raw_path)


def add_markdown_to_doc(md_text: str, doc: Document) -> None:
    image_re = re.compile(r"!\[(.*?)\]\((.*?)\)")
    ordered_re = re.compile(r"^\d+\.\s+(.*)$")

    for raw_line in md_text.splitlines():
        line = raw_line.rstrip()
        stripped = line.strip()

        if not stripped:
            doc.add_paragraph("")
            continue

        if stripped == "---":
            doc.add_paragraph("")
            continue

        if stripped.startswith("### "):
            doc.add_heading(stripped[4:].strip(), level=3)
            continue
        if stripped.startswith("## "):
            doc.add_heading(stripped[3:].strip(), level=2)
            continue
        if stripped.startswith("# "):
            doc.add_heading(stripped[2:].strip(), level=1)
            continue

        image_match = image_re.search(stripped)
        if image_match:
            alt_text = image_match.group(1).strip()
            image_path = markdown_image_path(image_match.group(2))
            if image_path.exists():
                doc.add_picture(str(image_path), width=Inches(6.5))
                if alt_text:
                    cap = doc.add_paragraph(alt_text)
                    cap.style = "Caption"
            else:
                doc.add_paragraph(f"[Missing image: {image_path}]")
            continue

        ordered_match = ordered_re.match(stripped)
        if ordered_match:
            doc.add_paragraph(ordered_match.group(1).strip(), style="List Number")
            continue

        if stripped.startswith("- "):
            doc.add_paragraph(stripped[2:].strip(), style="List Bullet")
            continue

        doc.add_paragraph(stripped)


def main() -> None:
    if not MD_PATH.exists():
        raise FileNotFoundError(f"Missing source markdown file: {MD_PATH}")

    document = Document()
    markdown_text = MD_PATH.read_text(encoding="utf-8")
    add_markdown_to_doc(markdown_text, document)
    document.save(str(OUT_PATH))
    print(f"Wrote: {OUT_PATH}")


if __name__ == "__main__":
    main()
